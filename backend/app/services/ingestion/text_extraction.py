"""Text extraction service for various document formats."""

from pathlib import Path
from typing import Tuple


class TextExtractionService:
    """Service for extracting text from various document formats."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"}

    def extract_text(self, file_path: str | Path) -> Tuple[str, str | None]:
        """
        Extract text from a document file.

        Args:
            file_path: Path to the document file

        Returns:
            Tuple of (extracted_text, error_message)
        """
        path = Path(file_path)

        if not path.exists():
            return "", f"File not found: {file_path}"

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            return "", f"Unsupported file type: {extension}"

        try:
            if extension == ".pdf":
                return self._extract_from_pdf(path), None
            elif extension in {".docx", ".doc"}:
                return self._extract_from_docx(path), None
            elif extension in {".txt", ".md"}:
                return self._extract_from_text(path), None
            else:
                return "", f"Unsupported file type: {extension}"
        except Exception as e:
            return "", f"Extraction error: {str(e)}"

    def _extract_from_pdf(self, path: Path) -> str:
        """Extract text from PDF file."""
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts)

    def _extract_from_docx(self, path: Path) -> str:
        """Extract text from DOCX file."""
        from docx import Document

        doc = Document(str(path))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        return "\n\n".join(text_parts)

    def _extract_from_text(self, path: Path) -> str:
        """Extract text from plain text or markdown file."""
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    def extract_from_bytes(self, content: bytes, filename: str) -> Tuple[str, str | None]:
        """
        Extract text from file content bytes.

        Args:
            content: File content as bytes
            filename: Original filename (for extension detection)

        Returns:
            Tuple of (extracted_text, error_message)
        """
        import tempfile
        import os

        extension = Path(filename).suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            return "", f"Unsupported file type: {extension}"

        # Write to temp file and extract
        with tempfile.NamedTemporaryFile(delete=False, suffix=extension) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            return self.extract_text(tmp_path)
        finally:
            os.unlink(tmp_path)
